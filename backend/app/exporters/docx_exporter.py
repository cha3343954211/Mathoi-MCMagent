"""Markdown -> docx 导出器。

策略：
1. 优先使用 pandoc（若系统已安装）→ 高保真 LaTeX 公式渲染；
2. 退化到 python-docx 简易渲染：标题、段落、图片、列表。
"""
from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

from ..core.logging import logger


def export_paper(md_path: Path, out_path: Path) -> Path:
    if not md_path.exists():
        raise FileNotFoundError(md_path)

    # 优先 pandoc
    if shutil.which("pandoc"):
        try:
            subprocess.run(
                ["pandoc", str(md_path), "-o", str(out_path), "--standalone"],
                check=True,
                cwd=str(md_path.parent),
                capture_output=True,
            )
            logger.info("pandoc export ok: {}", out_path)
            return out_path
        except subprocess.CalledProcessError as e:
            logger.warning("pandoc failed, fallback: {}", e.stderr.decode(errors="ignore"))

    return _docx_fallback(md_path, out_path)


def _docx_fallback(md_path: Path, out_path: Path) -> Path:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    work_dir = md_path.parent
    text = md_path.read_text(encoding="utf-8")

    for line in text.splitlines():
        s = line.rstrip()
        if not s.strip():
            doc.add_paragraph("")
            continue

        # 图片
        m = re.match(r"^!\[.*?\]\(([^)]+)\)\s*$", s)
        if m:
            img = work_dir / m.group(1)
            if img.exists():
                try:
                    doc.add_picture(str(img), width=Inches(5.5))
                    continue
                except Exception:
                    pass
            doc.add_paragraph(s)
            continue

        # 标题
        if s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("#### "):
            doc.add_heading(s[5:], level=4)
        elif s.startswith(("- ", "* ")):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", s):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", s), style="List Number")
        else:
            doc.add_paragraph(s)

    doc.save(str(out_path))
    logger.info("docx fallback export: {}", out_path)
    return out_path
